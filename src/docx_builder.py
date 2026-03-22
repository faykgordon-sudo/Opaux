"""
docx_builder.py -- Generate .docx CV files in American, German, or Europass formats.
"""

import os
from typing import Any

from rich.console import Console

console = Console()

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _require_docx() -> None:
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for CV generation.\n"
            "Install with: pip install python-docx"
        )


def _set_margins(doc: Any, margin_inches: float) -> None:
    """Set all page margins on the document."""
    for section in doc.sections:
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)


def _add_heading(doc: Any, text: str, level: int = 1, font_name: str = "Arial",
                 font_size: int = 13, bold: bool = True, color: tuple | None = None,
                 space_before: int = 12, space_after: int = 4) -> Any:
    """Add a styled section heading to the document."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Add bottom border for section heading
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def _add_paragraph(doc: Any, text: str, font_name: str = "Arial",
                   font_size: int = 11, bold: bool = False,
                   italic: bool = False, space_before: int = 0,
                   space_after: int = 2) -> Any:
    """Add a styled paragraph to the document."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    return para


def _add_bullet(doc: Any, text: str, font_name: str = "Arial",
                font_size: int = 11) -> Any:
    """Add a bullet point paragraph."""
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    return para


def _format_date_range(start: str, end: str, date_fmt: str = "MMM YYYY") -> str:
    """Format start/end dates into a range string."""
    end_str = "Present" if str(end).lower() in ("present", "current", "") else end
    return f"{start} - {end_str}"


def _get_profile_from_content(content: dict) -> dict:
    """Extract profile/personal data from tailored content."""
    return content.get("profile", {}) or content.get("personal", {})


# ---------------------------------------------------------------------------
# American format
# ---------------------------------------------------------------------------

def build_american(content: dict, format_config: dict) -> Any:
    """
    Build an American-style CV document.
    Single column, Arial 11pt, ATS-optimized, 2 pages max.
    """
    _require_docx()

    doc = Document()
    rules = format_config.get("rules", {})
    headers = format_config.get("headers", {})
    font = rules.get("font", "Arial")
    font_size = int(rules.get("font_size", 11))
    margin = float(rules.get("margins_inches", 0.75))
    max_bullets = int(rules.get("max_bullets_per_job", 5))

    _set_margins(doc, margin)

    profile = content.get("profile", {})
    personal = profile.get("personal", profile)

    # --- Header: Name ---
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(4)
    name_run = name_para.add_run(personal.get("name", "Your Name"))
    name_run.font.name = font
    name_run.font.size = Pt(18)
    name_run.font.bold = True

    # Contact line
    contact_parts = []
    for field in ("email", "phone", "location", "linkedin"):
        val = personal.get(field, "")
        if val:
            contact_parts.append(val)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_before = Pt(0)
    contact_para.paragraph_format.space_after = Pt(8)
    contact_run = contact_para.add_run("  |  ".join(contact_parts))
    contact_run.font.name = font
    contact_run.font.size = Pt(9)

    # --- Professional Summary ---
    summary_text = content.get("tailored_summary") or profile.get("summary", "")
    if summary_text:
        _add_heading(doc, headers.get("summary", "Professional Summary"),
                     font_name=font, font_size=12)
        _add_paragraph(doc, summary_text.strip(), font_name=font,
                       font_size=font_size, space_after=4)

    # --- Work Experience ---
    experience = profile.get("experience", [])
    tailored_bullets = content.get("tailored_bullets", {})

    if experience:
        _add_heading(doc, headers.get("experience", "Work Experience"),
                     font_name=font, font_size=12)

        for job in experience:
            title = job.get("title", "")
            company = job.get("company", "")
            location = job.get("location", "")
            start = job.get("start", "")
            end = job.get("end", "present")
            date_range = _format_date_range(start, end)

            # Job header row (Title | Company -- Location | Date)
            job_para = doc.add_paragraph()
            job_para.paragraph_format.space_before = Pt(6)
            job_para.paragraph_format.space_after = Pt(1)

            title_run = job_para.add_run(title)
            title_run.font.name = font
            title_run.font.size = Pt(font_size)
            title_run.font.bold = True

            sep_run = job_para.add_run(f"  -  {company}")
            if location:
                sep_run = job_para.add_run(f"  -  {company}, {location}")
            sep_run.font.name = font
            sep_run.font.size = Pt(font_size)
            sep_run.font.bold = False

            date_run = job_para.add_run(f"  |  {date_range}")
            date_run.font.name = font
            date_run.font.size = Pt(font_size)
            date_run.font.italic = True

            # Bullets: prefer tailored bullets, fall back to original
            bullets = tailored_bullets.get(title, job.get("bullets", []))
            for bullet in bullets[:max_bullets]:
                _add_bullet(doc, bullet, font_name=font, font_size=font_size)

    # --- Skills ---
    skills_data = profile.get("skills", {})
    skills_to_highlight = content.get("skills_to_highlight", [])

    if skills_data or skills_to_highlight:
        _add_heading(doc, headers.get("skills", "Skills"),
                     font_name=font, font_size=12)

        all_skills: dict = {}
        if isinstance(skills_data, dict):
            all_skills = skills_data
        elif isinstance(skills_data, list):
            all_skills = {"Skills": skills_data}

        for category, skill_list in all_skills.items():
            if not skill_list:
                continue
            cat_name = str(category).replace("_", " ").title()
            skill_str = ", ".join(str(s) for s in skill_list)
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(1)
            cat_run = para.add_run(f"{cat_name}: ")
            cat_run.font.name = font
            cat_run.font.size = Pt(font_size)
            cat_run.font.bold = True
            val_run = para.add_run(skill_str)
            val_run.font.name = font
            val_run.font.size = Pt(font_size)

        if skills_to_highlight:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(1)
            hl_run = para.add_run("Key Skills: ")
            hl_run.font.name = font
            hl_run.font.size = Pt(font_size)
            hl_run.font.bold = True
            val_run = para.add_run(", ".join(skills_to_highlight))
            val_run.font.name = font
            val_run.font.size = Pt(font_size)

    # --- Education ---
    education = profile.get("education", [])
    if education:
        _add_heading(doc, headers.get("education", "Education"),
                     font_name=font, font_size=12)
        for edu in education:
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_before = Pt(4)
            edu_para.paragraph_format.space_after = Pt(1)

            degree_run = edu_para.add_run(edu.get("degree", ""))
            degree_run.font.name = font
            degree_run.font.size = Pt(font_size)
            degree_run.font.bold = True

            institution = edu.get("institution", "")
            year = edu.get("year", "")
            grade = edu.get("grade", "")
            detail = f"  -  {institution}"
            if year:
                detail += f"  |  {year}"
            if grade:
                detail += f"  |  {grade}"
            det_run = edu_para.add_run(detail)
            det_run.font.name = font
            det_run.font.size = Pt(font_size)

    # --- Certifications ---
    certifications = profile.get("certifications", [])
    if certifications:
        _add_heading(doc, headers.get("certifications", "Certifications"),
                     font_name=font, font_size=12)
        for cert in certifications:
            cert_para = doc.add_paragraph()
            cert_para.paragraph_format.space_before = Pt(2)
            cert_para.paragraph_format.space_after = Pt(1)
            cert_name = cert.get("name", "")
            issuer = cert.get("issuer", "")
            date = cert.get("date", "")
            cert_run = cert_para.add_run(cert_name)
            cert_run.font.name = font
            cert_run.font.size = Pt(font_size)
            cert_run.font.bold = True
            if issuer or date:
                detail = ""
                if issuer:
                    detail += f"  -  {issuer}"
                if date:
                    detail += f"  |  {date}"
                det_run = cert_para.add_run(detail)
                det_run.font.name = font
                det_run.font.size = Pt(font_size)

    return doc


# ---------------------------------------------------------------------------
# German format
# ---------------------------------------------------------------------------

def _add_two_col_row(doc: Any, label: str, value: str, font: str = "Arial",
                     font_size: int = 11) -> None:
    """Add a label-value row using a two-column table."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(5.0)

    # Remove borders
    for cell in table.rows[0].cells:
        for border_name in ("top", "bottom", "left", "right", "insideH", "insideV"):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tcBorders.append(border)
            tcPr.append(tcBorders)

    label_cell = table.rows[0].cells[0]
    label_para = label_cell.paragraphs[0]
    label_run = label_para.add_run(label)
    label_run.font.name = font
    label_run.font.size = Pt(font_size)
    label_run.font.bold = True

    value_cell = table.rows[0].cells[1]
    value_para = value_cell.paragraphs[0]
    value_run = value_para.add_run(value)
    value_run.font.name = font
    value_run.font.size = Pt(font_size)


def _add_photo_textbox(doc: Any, placeholder: str = "Foto hier einfügen") -> None:
    """
    Add a photo placeholder text box in the top-right corner using drawing XML.
    This uses inline XML to create a floating text box.
    """
    # We add a paragraph with an inline shape (text box) via OOXML
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

    # Build the drawing XML for a floating text box
    drawing_xml = f"""
    <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:rPr/>
      <mc:AlternateContent xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">
        <mc:Choice xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
                   xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                   Requires="wps">
          <w:drawing>
            <wp:anchor distT="0" distB="0" distL="114300" distR="114300"
                       simplePos="0" relativeHeight="251658240" behindDoc="0"
                       locked="0" layoutInCell="1" allowOverlap="1">
              <wp:simplePos x="0" y="0"/>
              <wp:positionH relativeFrom="margin">
                <wp:align>right</wp:align>
              </wp:positionH>
              <wp:positionV relativeFrom="margin">
                <wp:posOffset>0</wp:posOffset>
              </wp:positionV>
              <wp:extent cx="1143000" cy="1371600"/>
              <wp:effectExtent l="0" t="0" r="0" b="0"/>
              <wp:wrapNone/>
              <wp:docPr id="1" name="PhotoBox"/>
              <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
                <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                  <wps:wsp>
                    <wps:cNvSpPr><a:spLocks noChangeArrowheads="1"/></wps:cNvSpPr>
                    <wps:spPr>
                      <a:xfrm><a:off x="0" y="0"/><a:ext cx="1143000" cy="1371600"/></a:xfrm>
                      <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                      <a:ln w="12700"><a:solidFill><a:srgbClr val="000000"/></a:solidFill></a:ln>
                    </wps:spPr>
                    <wps:txbx>
                      <w:txbxContent>
                        <w:p>
                          <w:pPr><w:jc w:val="center"/></w:pPr>
                          <w:r><w:t>{placeholder}</w:t></w:r>
                        </w:p>
                      </w:txbxContent>
                    </wps:txbx>
                    <wps:bodyPr anchor="ctr"/>
                  </wps:wsp>
                </a:graphicData>
              </a:graphic>
            </wp:anchor>
          </w:drawing>
        </mc:Choice>
      </mc:AlternateContent>
    </w:r>
    """

    try:
        from lxml import etree
        drawing_element = etree.fromstring(drawing_xml.strip())
        para._p.append(drawing_element)
    except Exception:
        # Fallback: add a simple bracketed placeholder
        run = para.add_run(f"[{placeholder}]")
        run.font.size = Pt(9)
        run.font.italic = True
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def build_german(content: dict, format_config: dict) -> Any:
    """
    Build a German-style Lebenslauf document.
    Tabular layout, photo placeholder, personal data, signature block.
    """
    _require_docx()

    doc = Document()
    rules = format_config.get("rules", {})
    headers = format_config.get("headers", {})
    font = rules.get("font", "Arial")
    font_size = int(rules.get("font_size", 11))
    margin = float(rules.get("margins_inches", 1.0))
    photo_placeholder = rules.get("photo_placeholder", "Foto hier einfügen")
    signature_block = rules.get("signature_block", "Ort, Datum                    Unterschrift")

    _set_margins(doc, margin)

    profile = content.get("profile", {})
    personal = profile.get("personal", profile)
    tailored_bullets = content.get("tailored_bullets", {})

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Lebenslauf")
    title_run.font.name = font
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_para.paragraph_format.space_after = Pt(12)

    # Photo placeholder (top-right)
    if rules.get("include_photo", True):
        _add_photo_textbox(doc, photo_placeholder)

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(personal.get("name", ""))
    name_run.font.name = font
    name_run.font.size = Pt(14)
    name_run.font.bold = True
    name_para.paragraph_format.space_after = Pt(12)

    # --- Persönliche Daten (Personal Data) ---
    _add_heading(doc, headers.get("personal_data", "Persönliche Angaben"),
                 font_name=font, font_size=12)

    fields_to_show = [
        ("Adresse", personal.get("location", "")),
        ("Telefon", personal.get("phone", "")),
        ("E-Mail", personal.get("email", "")),
        ("LinkedIn", personal.get("linkedin", "")),
    ]
    if rules.get("include_dob", True):
        fields_to_show.append(("Geburtsdatum", personal.get("date_of_birth", "")))
    if rules.get("include_nationality", True):
        fields_to_show.append(("Nationalität", personal.get("nationality", "")))
    if rules.get("include_marital_status", True):
        fields_to_show.append(("Familienstand", personal.get("marital_status", "")))

    for label, value in fields_to_show:
        if value:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            lbl_run = para.add_run(f"{label}: ")
            lbl_run.font.name = font
            lbl_run.font.size = Pt(font_size)
            lbl_run.font.bold = True
            val_run = para.add_run(value)
            val_run.font.name = font
            val_run.font.size = Pt(font_size)

    # --- Berufserfahrung ---
    experience = profile.get("experience", [])
    if experience:
        _add_heading(doc, headers.get("experience", "Berufserfahrung"),
                     font_name=font, font_size=12)

        for job in experience:
            title_str = job.get("title", "")
            company = job.get("company", "")
            location_str = job.get("location", "")
            start = job.get("start", "")
            end = job.get("end", "heute")
            if str(end).lower() in ("present", "current"):
                end = "heute"

            # Two-column: date | title & company
            table = doc.add_table(rows=1, cols=2)
            table.columns[0].width = Inches(1.3)
            table.columns[1].width = Inches(5.0)

            date_cell = table.rows[0].cells[0]
            date_para = date_cell.paragraphs[0]
            date_para.paragraph_format.space_before = Pt(4)
            date_run = date_para.add_run(f"{start} -\n{end}")
            date_run.font.name = font
            date_run.font.size = Pt(font_size)
            date_run.font.bold = True

            content_cell = table.rows[0].cells[1]
            content_para = content_cell.paragraphs[0]
            content_para.paragraph_format.space_before = Pt(4)
            title_run = content_para.add_run(title_str)
            title_run.font.name = font
            title_run.font.size = Pt(font_size)
            title_run.font.bold = True
            company_run = content_para.add_run(f"\n{company}")
            if location_str:
                company_run = content_para.add_run(f"\n{company}, {location_str}")
            company_run.font.name = font
            company_run.font.size = Pt(font_size)

            bullets = tailored_bullets.get(title_str, job.get("bullets", []))
            for bullet in bullets[:5]:
                b_para = content_cell.add_paragraph(style="List Bullet")
                b_run = b_para.add_run(bullet)
                b_run.font.name = font
                b_run.font.size = Pt(font_size)

    # --- Ausbildung ---
    education = profile.get("education", [])
    if education:
        _add_heading(doc, headers.get("education", "Ausbildung"),
                     font_name=font, font_size=12)

        for edu in education:
            table = doc.add_table(rows=1, cols=2)
            table.columns[0].width = Inches(1.3)
            table.columns[1].width = Inches(5.0)

            date_cell = table.rows[0].cells[0]
            date_run = date_cell.paragraphs[0].add_run(edu.get("year", ""))
            date_run.font.name = font
            date_run.font.size = Pt(font_size)
            date_run.font.bold = True
            date_cell.paragraphs[0].paragraph_format.space_before = Pt(4)

            content_cell = table.rows[0].cells[1]
            content_para = content_cell.paragraphs[0]
            content_para.paragraph_format.space_before = Pt(4)
            deg_run = content_para.add_run(edu.get("degree", ""))
            deg_run.font.name = font
            deg_run.font.size = Pt(font_size)
            deg_run.font.bold = True
            inst = edu.get("institution", "")
            if inst:
                inst_run = content_para.add_run(f"\n{inst}")
                inst_run.font.name = font
                inst_run.font.size = Pt(font_size)

    # --- Kenntnisse (Skills) ---
    skills_data = profile.get("skills", {})
    if skills_data:
        _add_heading(doc, headers.get("skills", "Kenntnisse"),
                     font_name=font, font_size=12)

        if isinstance(skills_data, dict):
            for category, skill_list in skills_data.items():
                if not skill_list:
                    continue
                cat_display = {
                    "languages": "Programmiersprachen",
                    "frameworks": "Frameworks",
                    "databases": "Datenbanken",
                    "cloud": "Cloud / DevOps",
                    "tools": "Tools",
                    "soft": "Soziale Kompetenzen",
                }.get(str(category).lower(), str(category).title())

                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(1)
                cat_run = para.add_run(f"{cat_display}: ")
                cat_run.font.name = font
                cat_run.font.size = Pt(font_size)
                cat_run.font.bold = True
                val_run = para.add_run(", ".join(str(s) for s in skill_list))
                val_run.font.name = font
                val_run.font.size = Pt(font_size)

    # --- Weiterbildung (Certifications) ---
    certifications = profile.get("certifications", [])
    if certifications:
        _add_heading(doc, headers.get("certifications", "Weiterbildung"),
                     font_name=font, font_size=12)
        for cert in certifications:
            cert_para = doc.add_paragraph()
            cert_para.paragraph_format.space_before = Pt(2)
            cert_para.paragraph_format.space_after = Pt(1)
            name_run = cert_para.add_run(cert.get("name", ""))
            name_run.font.name = font
            name_run.font.size = Pt(font_size)
            name_run.font.bold = True
            detail = ""
            if cert.get("issuer"):
                detail += f"  -  {cert['issuer']}"
            if cert.get("date"):
                detail += f"  |  {cert['date']}"
            if detail:
                det_run = cert_para.add_run(detail)
                det_run.font.name = font
                det_run.font.size = Pt(font_size)

    # --- Sprachkenntnisse ---
    languages = profile.get("languages", [])
    if languages:
        _add_heading(doc, headers.get("languages", "Sprachkenntnisse"),
                     font_name=font, font_size=12)
        for lang in languages:
            lang_para = doc.add_paragraph()
            lang_para.paragraph_format.space_before = Pt(2)
            lang_para.paragraph_format.space_after = Pt(1)
            lang_run = lang_para.add_run(lang.get("language", ""))
            lang_run.font.name = font
            lang_run.font.size = Pt(font_size)
            lang_run.font.bold = True
            level = lang.get("level", "")
            cefr = lang.get("cefr", "")
            detail = f": {level}"
            if cefr:
                detail += f" ({cefr})"
            det_run = lang_para.add_run(detail)
            det_run.font.name = font
            det_run.font.size = Pt(font_size)

    # --- Signature block ---
    if rules.get("include_signature", True):
        for _ in range(3):
            doc.add_paragraph()
        sig_para = doc.add_paragraph()
        sig_para.paragraph_format.space_before = Pt(24)
        sig_run = sig_para.add_run(signature_block)
        sig_run.font.name = font
        sig_run.font.size = Pt(font_size)

    return doc


# ---------------------------------------------------------------------------
# Europass format
# ---------------------------------------------------------------------------

def _add_europass_label_value(doc: Any, label: str, value: str,
                               font: str = "Arial", font_size: int = 11) -> None:
    """Add a label-content row in Europass style."""
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(5.0)

    label_cell = table.rows[0].cells[0]
    label_para = label_cell.paragraphs[0]
    label_para.paragraph_format.space_before = Pt(2)
    label_run = label_para.add_run(label.upper())
    label_run.font.name = font
    label_run.font.size = Pt(9)
    label_run.font.bold = True
    label_run.font.color.rgb = RGBColor(0x00, 0x4B, 0x87)  # Europass blue

    value_cell = table.rows[0].cells[1]
    value_para = value_cell.paragraphs[0]
    value_para.paragraph_format.space_before = Pt(2)
    value_run = value_para.add_run(value)
    value_run.font.name = font
    value_run.font.size = Pt(font_size)


def _add_cefr_grid(doc: Any, languages: list, cefr_levels: list, cefr_skills: list,
                   font: str = "Arial") -> None:
    """Add the CEFR language competence grid table."""
    doc.add_paragraph()
    grid_heading = doc.add_paragraph()
    grid_run = grid_heading.add_run("SELF-ASSESSMENT OF LANGUAGE COMPETENCES")
    grid_run.font.name = font
    grid_run.font.size = Pt(9)
    grid_run.font.bold = True
    grid_run.font.color.rgb = RGBColor(0x00, 0x4B, 0x87)

    # Header row: skill columns grouped under pairs
    num_skills = len(cefr_skills)
    table = doc.add_table(rows=2 + len(languages), cols=1 + num_skills)

    # Row 0: group headers
    group_headers = [
        ("UNDERSTANDING", 2),
        ("SPEAKING", 2),
        ("WRITING", 1),
    ]
    # Row 0, col 0: empty
    table.rows[0].cells[0].text = ""

    col_idx = 1
    for group_name, span in group_headers:
        cell = table.rows[0].cells[col_idx]
        cell.text = group_name
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        col_idx += span

    # Row 1: skill names
    table.rows[1].cells[0].text = "Language"
    table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
    table.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(8)

    for i, skill in enumerate(cefr_skills):
        cell = table.rows[1].cells[i + 1]
        cell.text = skill
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)

    # Data rows
    for j, lang in enumerate(languages):
        row = table.rows[2 + j]
        row.cells[0].text = lang.get("language", "")
        cefr = lang.get("cefr", "B1")
        for k in range(num_skills):
            row.cells[k + 1].text = cefr
            row.cells[k + 1].paragraphs[0].runs[0].font.size = Pt(9)

    table.style = "Table Grid"


def _add_digcomp_grid(doc: Any, digcomp_areas: list, digcomp_levels: list,
                      digital_skills: dict, font: str = "Arial") -> None:
    """Add the DIGCOMP digital competence grid table."""
    doc.add_paragraph()
    grid_heading = doc.add_paragraph()
    grid_run = grid_heading.add_run("DIGITAL COMPETENCES (DIGCOMP FRAMEWORK)")
    grid_run.font.name = font
    grid_run.font.size = Pt(9)
    grid_run.font.bold = True
    grid_run.font.color.rgb = RGBColor(0x00, 0x4B, 0x87)

    table = doc.add_table(rows=1 + len(digcomp_areas), cols=2)

    # Header
    table.rows[0].cells[0].text = "Competence Area"
    table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    table.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(8)
    table.rows[0].cells[1].text = "Level"
    table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
    table.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(8)

    for i, area in enumerate(digcomp_areas):
        table.rows[i + 1].cells[0].text = area
        table.rows[i + 1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        level = digital_skills.get(area, "Intermediate")
        table.rows[i + 1].cells[1].text = level
        table.rows[i + 1].cells[1].paragraphs[0].runs[0].font.size = Pt(9)

    table.style = "Table Grid"


def build_europass(content: dict, format_config: dict) -> Any:
    """
    Build a Europass-format CV document.
    Label-content layout, CEFR grid, DIGCOMP grid, EQF education labels.
    """
    _require_docx()

    doc = Document()
    rules = format_config.get("rules", {})
    headers = format_config.get("headers", {})
    extras = format_config.get("extras", format_config)
    font = rules.get("font", "Arial")
    font_size = int(rules.get("font_size", 11))
    margin = float(rules.get("margins_inches", 0.75))

    _set_margins(doc, margin)

    profile = content.get("profile", {})
    personal = profile.get("personal", profile)
    tailored_bullets = content.get("tailored_bullets", {})

cefr_skills = extras.get("cefr_skills", [
        "Listening", "Reading", "Spoken interaction", "Spoken production", "Writing"])
    ])
    digcomp_levels = extras.get("digcomp_levels", ["Foundation", "Intermediate", "Advanced", "Highly Specialised"])

    # Europass header bar
    header_para = doc.add_paragraph()
    header_para.paragraph_format.space_before = Pt(0)
    header_para.paragraph_format.space_after = Pt(4)
    header_run = header_para.add_run("Curriculum Vitae")
    header_run.font.name = font
    header_run.font.size = Pt(20)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(0x00, 0x4B, 0x87)

    # Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(personal.get("name", ""))
    name_run.font.name = font
    name_run.font.size = Pt(16)
    name_run.font.bold = True
    name_para.paragraph_format.space_after = Pt(8)

    # --- Personal Information ---
    _add_heading(doc, headers.get("contact", "Personal Information"),
                 font_name=font, font_size=11, color=(0x00, 0x4B, 0x87))

    contact_fields = [
        ("Address", personal.get("location", "")),
        ("Telephone", personal.get("phone", "")),
        ("Email", personal.get("email", "")),
        ("LinkedIn", personal.get("linkedin", "")),
        ("GitHub", personal.get("github", "")),
    ]
    if rules.get("include_dob", True):
        contact_fields.append(("Date of birth", personal.get("date_of_birth", "")))
    if rules.get("include_nationality", True):
        contact_fields.append(("Nationality", personal.get("nationality", "")))

    for label, value in contact_fields:
        if value:
            _add_europass_label_value(doc, label, value, font, font_size)

    # --- Desired Position ---
    desired = content.get("tailored_summary", profile.get("summary", ""))
    if desired:
        doc.add_paragraph()
        _add_heading(doc, headers.get("desired_position", "Desired Position / Occupational Field"),
                     font_name=font, font_size=11, color=(0x00, 0x4B, 0x87))
        _add_paragraph(doc, desired.strip(), font_name=font, font_size=font_size)

    # --- Work Experience ---
    experience = profile.get("experience", [])
    if experience:
        doc.add_paragraph()
        _add_heading(doc, headers.get("experience", "Work Experience"),
                     font_name=font, font_size=11, color=(0x00, 0x4B, 0x87))

        for job in experience:
            title_str = job.get("title", "")
            company = job.get("company", "")
            location_str = job.get("location", "")
            start = job.get("start", "")
            end = job.get("end", "Present")

           employer_str = f"{company}, {location_str}" if location_str else company
            _add_europass_label_value(doc, "Employer", employer_str, font, font_size)
            _add_europass_label_value(doc, "Occupation", title_str, font, font_size)
            _add_europass_label_value(doc, "Employer", f"{company}, {location_str}" if location_str else company, font, font_size)

            bullets = tailored_bullets.get(title_str, job.get("bullets", []))
            if bullets:
                for bullet in bullets[:5]:
                    _add_bullet(doc, bullet, font_name=font, font_size=font_size)

            doc.add_paragraph()

    # --- Education and Training ---
    education = profile.get("education", [])
    eqf_map = {"bachelor": "EQF 6", "master": "EQF 7", "phd": "EQF 8", "doctorate": "EQF 8",
                "bsc": "EQF 6", "msc": "EQF 7", "ba ": "EQF 6", "ma ": "EQF 7"}

    if education:
        doc.add_paragraph()
        _add_heading(doc, headers.get("education", "Education and Training"),
                     font_name=font, font_size=11, color=(0x00, 0x4B, 0x87))

        for edu in education:
            degree = edu.get("degree", "")
            institution = edu.get("institution", "")
            year = edu.get("year", "")
            grade = edu.get("grade", "")

            _add_europass_label_value(doc, "Dates", year, font, font_size)
            _add_europass_label_value(doc, "Title", degree, font, font_size)
            _add_europass_label_value(doc, "Organisation", institution, font, font_size)
            if grade:
                _add_europass_label_value(doc, "Grade", grade, font, font_size)

            # EQF level
            if rules.get("use_eqf_levels", True):
                degree_lower = degree.lower()
                eqf_label = "EQF 6"
                for keyword, label in eqf_map.items():
                    if keyword in degree_lower:
                        eqf_label = label
                        break
                _add_europass_label_value(doc, "EQF Level", eqf_label, font, font_size)

            doc.add_paragraph()

    # --- Personal Skills ---
    skills_data = profile.get("skills", {})
    if skills_data:
        doc.add_paragraph()
        _add_heading(doc, headers.get("skills", "Personal Skills"),
                     font_name=font, font_size=11, color=(0x00, 0x4B, 0x87))

        if isinstance(skills_data, dict):
            for category, skill_list in skills_data.items():
                if not skill_list:
                    continue
                _add_europass_label_value(
                    doc,
                    str(category).title(),
                    ", ".join(str(s) for s in skill_list),
                    font, font_size,
                )

    # --- Mother Tongue / Other Languages + CEFR Grid ---
    languages = profile.get("languages", [])
    if languages:
        doc.add_paragraph()
        _add_heading(doc, headers.get("languages", "Mother Tongue / Other Languages"),
                     font_name=font, font_size=11, color=(0x00, 0x4B, 0x87))

        mother_tongue = [l for l in languages if l.get("cefr", "").upper() in ("C2", "NATIVE")]
        other_langs = [l for l in languages if l not in mother_tongue]

        for lang in mother_tongue:
            _add_europass_label_value(doc, "Mother tongue(s)", lang.get("language", ""), font, font_size)

        if rules.get("use_cefr_grid", True) and other_langs:
            _add_cefr_grid(doc, other_langs, cefr_levels, cefr_skills, font)

    # --- Digital Competence + DIGCOMP Grid ---
    digital_skills = content.get("digital_skills", {})
    if rules.get("use_digcomp_grid", True):
        doc.add_paragraph()
        _add_heading(doc, headers.get("digital_competence", "Digital Competence"),
                     font_name=font, font_size=11, color=(0x00, 0x4B, 0x87))
        _add_digcomp_grid(doc, digcomp_areas, digcomp_levels, digital_skills, font)

    return doc


# ---------------------------------------------------------------------------
# Public interface
# ---------------------------------------------------------------------------

def save_document(doc: Any, output_path: str) -> str:
    """Save a python-docx Document to the given path. Returns the path."""
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return output_path


def build_cv(
    tailored_content: dict,
    format_config: dict,
    format_name: str,
    language: str,
    output_path: str,
) -> str:
    """
    Build a CV document for the given format, save it, and return the output path.

    Args:
        tailored_content: Dict from tailoring.py (includes profile, tailored_bullets, etc.)
        format_config: Dict loaded from cv_formats.py
        format_name: 'american', 'german', or 'europass'
        language: ISO 639-1 language code (for future use in doc metadata)
        output_path: Full path where the .docx should be saved

    Returns:
        The output_path string.
    """
    builders = {
        "american": build_american,
        "german": build_german,
        "europass": build_europass,
    }

    builder = builders.get(format_name)
    if not builder:
        raise ValueError(
            f"Unknown format '{format_name}'. Valid: {', '.join(builders.keys())}"
        )

    console.print(f"[blue]Building {format_name} CV...[/blue]")
    doc = builder(tailored_content, format_config)
    path = save_document(doc, output_path)
    console.print(f"[green]CV saved to:[/green] {path}")
    return path
