"""
cv_parser.py -- Parse an uploaded CV (PDF or DOCX) into a structured profile
                using Claude, and suggest job search terms.
"""

import json
import os
import time
from pathlib import Path
from typing import Any

import yaml
from rich.console import Console

console = Console()

CV_PARSE_PROMPT = """\
You are an expert CV/resume parser. Extract all information from this CV and return \
it as structured YAML.

CV TEXT:
{cv_text}

Extract the following and return ONLY valid JSON (no markdown, no explanation):

{{
  "personal": {{
    "name": "<full name>",
    "email": "<email>",
    "phone": "<phone number>",
    "location": "<city, country>",
    "linkedin": "<linkedin url or handle, or empty>",
    "github": "<github url or handle, or empty>",
    "date_of_birth": "<YYYY-MM-DD or empty>",
    "nationality": "<nationality or empty>",
    "marital_status": "<Single/Married/etc or empty>"
  }},
  "summary": "<2-4 sentence professional summary capturing the candidate's \
core identity, top skills, and career focus. Write it in first/third person \
to match the CV tone.>",
  "experience": [
    {{
      "title": "<job title>",
      "company": "<company name>",
      "location": "<city, country>",
      "start": "<YYYY-MM>",
      "end": "<YYYY-MM or present>",
      "bullets": ["<achievement/responsibility>", ...],
      "skills": ["<skill used in this role>", ...]
    }}
  ],
  "education": [
    {{
      "degree": "<degree name>",
      "institution": "<institution name>",
      "location": "<city, country>",
      "year": "<graduation year or expected year>",
      "grade": "<grade/GPA or empty>"
    }}
  ],
  "skills": {{
    "languages": ["<programming languages>"],
    "frameworks": ["<frameworks/libraries>"],
    "databases": ["<databases>"],
    "cloud": ["<cloud platforms>"],
    "tools": ["<software tools, ERP systems, platforms>"],
    "soft": ["<soft skills>"]
  }},
  "certifications": [
    {{
      "name": "<certification name>",
      "issuer": "<issuing body>",
      "date": "<YYYY-MM or empty>"
    }}
  ],
  "languages": [
    {{
      "language": "<language name>",
      "level": "<fluency level>",
      "cefr": "<CEFR level e.g. B2 or empty>"
    }}
  ],
  "suggested_search_terms": [
    "<job title that matches this candidate's background>",
    "<another relevant job title>",
    "<another relevant job title>"
  ],
  "suggested_location": "<primary job search location based on CV>"
}}

Rules:
- Extract ALL work experience entries, even short contracts
- For bullets: extract verbatim if possible, otherwise summarise accurately
- For skills: only include what is explicitly mentioned or clearly implied
- suggested_search_terms: provide 3-5 specific job titles this person should search for, \
  based on their actual experience level and background (e.g. "Junior Logistics Manager", \
  "Supply Chain Coordinator", not generic terms)
- If a field is not present in the CV, use an empty string or empty list
- Return ONLY the JSON object, nothing else
"""


def extract_text_from_pdf(path: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is required: pip install pdfplumber")

    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    return "\n\n".join(text_parts)


def extract_text_from_docx(path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = Document(path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def extract_cv_text(file_path: str) -> str:
    """Extract text from a CV file (PDF or DOCX)."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Upload a PDF or DOCX file.")


def _call_claude(client: Any, prompt: str, max_retries: int = 2) -> str:
    """Call Claude API with retry."""
    last_error = None
    for attempt in range(max_retries):
        try:
            message = client.messages.create(
                model="claude-opus-4-6",
                max_tokens=4096,
                messages=[{"role": "user", "content": prompt}],
            )
            return message.content[0].text
        except Exception as exc:
            last_error = exc
            if attempt < max_retries - 1:
                time.sleep(2)
    raise RuntimeError(f"Claude API failed: {last_error}")


def _parse_json_response(text: str) -> dict:
    """Extract JSON from Claude response."""
    text = text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        text = "\n".join(lines[1:])
        if text.strip().endswith("```"):
            text = text.strip()[:-3].strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}") + 1
        if start >= 0 and end > start:
            return json.loads(text[start:end])
        raise ValueError("Could not parse JSON from Claude response")


def save_profile(profile_data: dict, profile_path: str = "config/profile.yaml") -> None:
    """Save the parsed profile data to config/profile.yaml."""
    Path(profile_path).parent.mkdir(parents=True, exist_ok=True)

    # Build clean profile dict matching the expected schema
    profile = {
        "personal": profile_data.get("personal", {}),
        "summary": profile_data.get("summary", ""),
        "experience": profile_data.get("experience", []),
        "education": profile_data.get("education", []),
        "skills": {
            "languages": profile_data.get("skills", {}).get("languages", []),
            "frameworks": profile_data.get("skills", {}).get("frameworks", []),
            "databases": profile_data.get("skills", {}).get("databases", []),
            "cloud": profile_data.get("skills", {}).get("cloud", []),
            "tools": profile_data.get("skills", {}).get("tools", []),
            "soft": profile_data.get("skills", {}).get("soft", []),
        },
        "certifications": profile_data.get("certifications", []),
        "languages": profile_data.get("languages", []),
    }

    with open(profile_path, "w", encoding="utf-8") as f:
        yaml.dump(profile, f, allow_unicode=True, default_flow_style=False, sort_keys=False)


def update_settings_from_cv(
    parsed: dict,
    settings_path: str = "config/settings.yaml",
) -> None:
    """Update settings.yaml with search terms and location suggested by the CV parser."""
    settings_path = Path(settings_path)
    config = {}
    if settings_path.exists():
        with open(settings_path, "r", encoding="utf-8") as f:
            config = yaml.safe_load(f) or {}

    suggestions = parsed.get("suggested_search_terms", [])
    location = parsed.get("suggested_location", "")

    config.setdefault("discovery", {})
    if suggestions:
        # Use first suggestion as primary search term
        config["discovery"]["search_term"] = suggestions[0]
        config["discovery"]["suggested_terms"] = suggestions
    if location:
        config["discovery"]["location"] = location

    settings_path.parent.mkdir(parents=True, exist_ok=True)
    with open(settings_path, "w", encoding="utf-8") as f:
        yaml.dump(config, f, allow_unicode=True, default_flow_style=False)


def run_cv_parser(
    api_key: str,
    file_path: str,
) -> dict:
    """
    Parse an uploaded CV file and extract structured profile data.

    Args:
        api_key: Claude API key
        file_path: Path to the uploaded CV (PDF or DOCX)

    Returns:
        Parsed profile dict with suggested_search_terms included
    """
    try:
        import anthropic
    except ImportError:
        raise ImportError("anthropic package is required: pip install anthropic")

    if not api_key or api_key == "YOUR_CLAUDE_API_KEY":
        raise ValueError("Claude API key is required to parse your CV.")

    client = anthropic.Anthropic(api_key=api_key)

    console.print(f"[blue]Extracting text from CV: {Path(file_path).name}[/blue]")
    cv_text = extract_cv_text(file_path)

    if not cv_text.strip():
        raise ValueError("Could not extract any text from the uploaded file. Try a different format.")

    console.print(f"[blue]Extracted {len(cv_text)} characters. Sending to Claude for analysis...[/blue]")

    prompt = CV_PARSE_PROMPT.format(cv_text=cv_text[:8000])
    response = _call_claude(client, prompt)
    parsed = _parse_json_response(response)

    console.print("[green]CV parsed successfully.[/green]")
    return parsed
