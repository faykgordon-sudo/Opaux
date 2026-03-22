"""
cover_letter.py -- AI-generated cover letters via Claude.
"""

import json
import os
import time
from pathlib import Path
from typing import Any

import yaml
from rich.console import Console

from src.database import get_connection, get_job_by_id, init_db, update_job

console = Console()

COVER_LETTER_PROMPT = """\
Write a professional cover letter for this job application.

Candidate: {candidate_summary}

Job: {title} at {company}
Location: {location}

Key Matches / Skills:
{key_matches}

Job Description (excerpt):
{description}

Instructions:
- 3 paragraphs only
- Opening paragraph: strong hook that shows genuine interest in the specific role and company, mention the role title
- Middle paragraph: 2-3 specific achievements from the candidate's background mapped to the job requirements
- Closing paragraph: confident call to action, express enthusiasm for discussing further
- Tone: professional but personable -- genuine, not robotic
- Do NOT start with "I am writing to apply"
- Do NOT use filler phrases like "I am passionate about", "I believe I would be a great fit"
- Be specific: reference real technologies, real metrics, real company context where possible
- Keep total length under 350 words

Return ONLY the cover letter text, no JSON, no headers.
"""


def _load_profile(config: dict | None = None) -> dict:
    """Load the user profile from config/profile.yaml or per-user path."""
    profile_path = Path((config or {}).get("_profile_path", "config/profile.yaml"))
    if not profile_path.exists():
        return {}
    with open(profile_path, encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


def _build_candidate_summary(profile: dict) -> str:
    """Build a concise candidate summary string from the profile."""
    personal = profile.get("personal", {})
    name = personal.get("name", "The candidate")
    summary = profile.get("summary", "")
    skills = profile.get("skills", {})

    top_skills: list[str] = []
    if isinstance(skills, dict):
        for category in ("languages", "frameworks", "cloud"):
            skill_list = skills.get(category, [])
            top_skills.extend(str(s) for s in skill_list[:3])

    candidate_str = f"Name: {name}\n"
    if summary:
        candidate_str += f"Summary: {summary.strip()}\n"
    if top_skills:
        candidate_str += f"Key Skills: {', '.join(top_skills[:8])}"

    return candidate_str


def _call_claude(client: Any, prompt: str, max_retries: int = 2) -> str:
    """Call Claude API with retry on failure."""
    last_error = None
    for attempt in range(max_retries):
        try:
            message = client.messages.create(
                model="claude-opus-4-5",
                max_tokens=1024,
                messages=[{"role": "user", "content": prompt}],
            )
            return message.content[0].text
        except Exception as exc:
            last_error = exc
            if attempt < max_retries - 1:
                console.print(
                    f"[yellow]Claude API error (attempt {attempt + 1}): {exc}. Retrying...[/yellow]"
                )
                time.sleep(2)
    raise RuntimeError(f"Claude API failed after {max_retries} attempts: {last_error}")


def _save_cover_letter_docx(text: str, output_path: str, profile: dict, lang: str = "en") -> str:
    """Save the cover letter text as a .docx file."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    personal = profile.get("personal", {})

    # Sender info
    sender_lines = []
    if personal.get("name"):
        sender_lines.append(personal["name"])
    if personal.get("email"):
        sender_lines.append(personal["email"])
    if personal.get("phone"):
        sender_lines.append(personal["phone"])
    if personal.get("location"):
        sender_lines.append(personal["location"])

    for line in sender_lines:
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()  # blank line

    # Cover letter body
    paragraphs = [p.strip() for p in text.strip().split("\n\n") if p.strip()]
    for para_text in paragraphs:
        para = doc.add_paragraph()
        run = para.add_run(para_text)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(10)

    # Signature -- language-aware closing
    _closings = {
        "de": "Mit freundlichen Grüßen",
        "fr": "Cordialement",
        "es": "Atentamente",
        "pt": "Atenciosamente",
        "nl": "Met vriendelijke groet",
        "it": "Cordiali saluti",
        "pl": "Z poważaniem",
        "sv": "Med vänliga hälsningar",
        "da": "Med venlig hilsen",
        "no": "Med vennlig hilsen",
    }
    closing = _closings.get(lang, "Sincerely")
    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    sig_run = sig_para.add_run(f"{closing},\n\n{personal.get('name', '')}")
    sig_run.font.name = "Arial"
    sig_run.font.size = Pt(11)

    doc.save(output_path)
    return output_path


def run_cover_letter(config: dict, db_path: str, job_id: str, output_dir: str | None = None) -> str:
    """
    Generate a cover letter for a specific job and save as .docx.

    Args:
        config: App configuration dict
        db_path: Path to the SQLite database
        job_id: The job ID to generate a cover letter for

    Returns:
        Path to the saved cover letter .docx file
    """
    try:
        import anthropic
    except ImportError:
        raise ImportError("anthropic package is required. Install with: pip install anthropic")

    api_key = config.get("claude_api_key", "")
    if not api_key or api_key == "YOUR_CLAUDE_API_KEY":
        raise ValueError("claude_api_key is not set in config/settings.yaml")

    client = anthropic.Anthropic(api_key=api_key)

    # Load job from DB
    init_db(db_path)
    conn = get_connection(db_path)
    try:
        job = get_job_by_id(conn, job_id)
    finally:
        conn.close()

    if not job:
        raise ValueError(f"Job '{job_id}' not found in database.")

    title = job.get("title", "Unknown")
    company = job.get("company", "Unknown")
    location = job.get("location", "")
    description = job.get("description") or ""
    cv_lang = job.get("cv_lang") or "en"

    # Parse keywords
    keywords_raw = job.get("keywords_matched") or "[]"
    try:
        keywords = json.loads(keywords_raw)
        key_matches_str = "\n".join(f"- {k}" for k in keywords[:15])
    except (json.JSONDecodeError, TypeError):
        key_matches_str = keywords_raw[:500]

    console.print(f"[bold blue]Generating cover letter for:[/bold blue] {title} @ {company}")

    # Load profile
    profile = _load_profile(config)
    candidate_summary = _build_candidate_summary(profile)

    # Build prompt
    prompt = COVER_LETTER_PROMPT.format(
        candidate_summary=candidate_summary,
        title=title,
        company=company,
        location=location,
        key_matches=key_matches_str or "(no keywords extracted yet)",
        description=description[:2000],
    )

    # Call Claude
    lang_label = cv_lang.upper() if cv_lang != "en" else "English"
    console.print(f"[blue]Calling Claude to write cover letter (language: {lang_label})...[/blue]")

    # If non-English, instruct Claude to write directly in target language
    if cv_lang != "en":
        from src.translator import LANGUAGE_NAMES
        lang_name = LANGUAGE_NAMES.get(cv_lang, cv_lang)
        prompt = prompt.rstrip() + f"\n\nWrite the cover letter entirely in {lang_name}. " \
            f"Adapt salutations and closing phrases to {lang_name} professional conventions."

    cover_letter_text = _call_claude(client, prompt)

    # Determine output path using standardised naming convention
    from src.utils import cover_filename
    base_dir = output_dir or os.path.join(config.get("output", {}).get("dir", "output"), "cover_letters")
    os.makedirs(base_dir, exist_ok=True)
    fname = cover_filename(company, title, cv_lang) + ".docx"
    output_path = os.path.join(base_dir, fname)

    # Save as docx
    final_path = _save_cover_letter_docx(cover_letter_text, output_path, profile, lang=cv_lang)

    # Update DB
    conn = get_connection(db_path)
    try:
        update_job(conn, job_id, cover_letter_path=final_path)
    finally:
        conn.close()

    console.print(f"[bold green]Cover letter saved to:[/bold green] {final_path}")
    console.print("\n[dim]--- Preview (first paragraph) ---[/dim]")
    first_para = cover_letter_text.strip().split("\n\n")[0] if cover_letter_text else ""
    console.print(f"[italic]{first_para[:300]}...[/italic]")

    return final_path
